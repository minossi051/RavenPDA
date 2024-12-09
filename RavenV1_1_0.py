 #patch 1.1.0
#função para gerar mapa interativo dos pops foi implementada ++
#bug da função atualizar_dados() não executado foi resolvido
#função secreta ++

import pygame
import time
import random
import plotly.graph_objects as go
import plotly.express as px
from geopy.distance import geodesic
import openai
import pandas as pd
import sys
import os
import json
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from tkinter import Tk, Label, Button, Menu, filedialog, messagebox, Toplevel, Text, Scrollbar, Frame, BOTH, RIGHT, LEFT, Y,X
from PIL import Image, ImageTk
import time as t
import webbrowser
import datetime
import getpass
import socket
import time
from docx import Document
from tkinter.scrolledtext import ScrolledText
# Configurações de rede
hostname = socket.gethostbyname




def get_ipv4():
    
    try:
        hostname = socket.gethostname()
        ip_address = socket.gethostbyname(hostname)
        return ip_address
    except Exception as e:
        print(f"Erro ao obter o endereço IPv4: {e}")
    return None

def check_network(expected_ssid, expected_ip):
    """
    Verifica se o SSID e o IPv4 correspondem aos valores esperados.
    """
    ip_address = get_ipv4()


    ip_match = (ip_address == expected_ip)

    if ip_match:
        return True, f"Conectado à rede com o IPv4 '{ip_address}'."

REDERS_ip = '172.28.7.158'
REDRS_ssid = 'REDERS'

status, message = check_network(REDRS_ssid,REDERS_ip)
#usuario
def get_computer_user():
    return getpass.getuser()

# horario
data_hora = datetime.datetime.now().strftime("%H,%M")

apr = ''
hora_atual = int(data_hora.split(',')[0]) * 60 + int(data_hora.split(',')[1])  # Convertendo para minutos

if hora_atual >= 5 * 60:  # 5:00
    apr = 'Bom dia'
if hora_atual >= 12 * 60 + 30:  # 12:30
    apr = 'Boa tarde'
if hora_atual >= 18 * 60 + 30:  # 18:30
    apr = 'Boa noite'


#usuarios pro-pae
def verificar_usuario_autorizado(user, nomes_autorizados):
    return user in nomes_autorizados

# Variáveis globais
nomes_autorizados = ['vicenzo-minossi', 'moises-brum', 'dario-wachholz', 'carlos-meneses']

#localizar a pasta do programa
def encontrar_pasta_ravenia():
    # Determina o diretório base dependendo se está rodando como script ou executável
    if getattr(sys, 'frozen', False):  # Se estiver no modo frozen (executável PyInstaller)
        caminho_base = sys._MEIPASS
    else:  # Se estiver rodando no ambiente de desenvolvimento
        caminho_base = os.path.dirname(os.path.abspath(__file__))

    # Subir na hierarquia até encontrar a pasta 'RavenIA' ou atingir a raiz
    while caminho_base != os.path.dirname(caminho_base):  # Enquanto não chegar à raiz
        caminho_ravenia = os.path.join(caminho_base, 'RavenIA')
        if os.path.isdir(caminho_ravenia):  # Verifica se o diretório 'RavenIA' existe
            return caminho_ravenia
        caminho_base = os.path.dirname(caminho_base)  # Subir um nível na hierarquia

    # Caso não encontre a pasta 'RavenIA', retorna None ou um erro, conforme a necessidade
    return None
#caminho da pasta global
def caminho_relativo(*caminhos):
    caminho_ravenia = encontrar_pasta_ravenia()
    
    if caminho_ravenia is None:
        raise FileNotFoundError("A pasta 'RavenIA' não foi encontrada.")
    
    return os.path.join(caminho_ravenia, *caminhos)
#problemas para executar o arquivo executável, não entendi por que o pyinstaller não acha o arquivo


# caminhos de arquivos
base_logs = caminho_relativo('atualizacoes.txt')
base_sobre = caminho_relativo('sobre.txt')
base_teste = caminho_relativo('patchlogs.txt')

#pngs
icon_atualizardados = caminho_relativo('data', 'pngs', 'atualizar.png')
icon_sobre = caminho_relativo('data', 'pngs', 'info.png')
icon_teste = caminho_relativo('data', 'pngs', 'news.png')
icon_consultaratt = caminho_relativo('data', 'pngs', 'search.png')
icon_acessardados = caminho_relativo('data', 'pngs', 'cloud.png')
icon_onedrive = caminho_relativo('data', 'pngs', 'onedrive.png')
icon_relatorio = caminho_relativo('data','pngs','writer.png')
icon_raven = caminho_relativo('data','pngs','raven.png')
icon_mapa = caminho_relativo('data','pngs','map.png')
fundo_game = caminho_relativo('data','pngs','mapars.png')


#logo do menuzão
logo_menu = caminho_relativo('data','pngs','logo.png')
# carregar caminho do OneDrive salvo ou usar um caminho padrão
config_path = caminho_relativo('bot_config.json')
if os.path.exists(config_path):
    with open(config_path, 'r') as config_file:
        config = json.load(config_file)
        pasta_onedrive = config.get('pasta_onedrive', caminho_relativo('OneDrive - Governo do Estado do Rio Grande do Sul', 'POPS', 'procergs-diop-dif-pir'))
else:
    pasta_onedrive = caminho_relativo('OneDrive - Governo do Estado do Rio Grande do Sul', 'POPS', 'procergs-diop-dif-pir')

# alterar ou encontrar pasta do one drive
def selecionar_pasta_onedrive():
    global pasta_onedrive
    nova_pasta = filedialog.askdirectory(title="Selecione a pasta do OneDrive")
    if nova_pasta:
        pasta_onedrive = nova_pasta
        with open(config_path, 'w') as config_file:
            json.dump({'pasta_onedrive': pasta_onedrive}, config_file)
        messagebox.showinfo("Sucesso", f"Pasta do OneDrive alterada para: {pasta_onedrive}")

# check user
admin = False
user = get_computer_user()
if user == 'vicenzo-minossi':
    admin = True
if not verificar_usuario_autorizado(user, nomes_autorizados):
    messagebox.showerror("Acesso negado", "Você não está autorizado a executar este programa.")
    exit()


# registro de log
def registrar_atualizacao(mensagem):
    operador_atual = get_computer_user()
    with open(base_logs, "a") as file:
        data_hora = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        file.write(f"{data_hora} - {mensagem} - Operador: {operador_atual}\n")

# selecionar arquivos
def selecionar_arquivo(titulo):
    caminho_arquivo = filedialog.askopenfilename(title=titulo, filetypes=[('Arquivos Excel', '*.xlsx'), ('Todos os arquivos', '*.*')])
    return caminho_arquivo

#relatorio word
def gerar_relatorio(nome_pop, endereco, municipio, data_atualizacao, detalhes):
    try:
        # Verifica se os campos obrigatórios estão preenchidos
        if not nome_pop or not endereco or not municipio or not data_atualizacao:
            messagebox.showwarning("Atenção", "Preencha todos os campos!")
            return
        
        # Cria o documento com base nas informações inseridas
        documento = Document()
        documento.add_heading("PROCERGS-DIOP-DIF-PIR", level=1)
        documento.add_paragraph(f"Nome POP: {nome_pop}")
        documento.add_paragraph(f"Endereço: {endereco}")
        documento.add_paragraph(f"Município: {municipio}")
        documento.add_paragraph(f"Data de atualização: {data_atualizacao}")
        documento.add_paragraph("Documento gerado automaticamente por RavenIA")

        # Descrição detalhada
        documento.add_heading("Descrição detalhada:", level=2)
        documento.add_paragraph(detalhes)

        # Salvar o documento
        caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if caminho:
            documento.save(caminho)
            messagebox.showinfo("Sucesso", f"Relatório salvo em: {caminho}")
        else:
            messagebox.showwarning("Aviso", "Salvamento cancelado.")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao gerar o relatório: {e}")

def abrir_janela_relatorio():
    janela_relatorio = Toplevel(root)
    janela_relatorio.title("Gerar Relatório")
    janela_relatorio.geometry("600x700")

    # Widgets para inserir informações
    Label(janela_relatorio, text="Nome POP:").pack(pady=5)
    entry_nome_pop = Text(janela_relatorio, height=1, width=50)
    entry_nome_pop.pack()

    Label(janela_relatorio, text="Endereço:").pack(pady=5)
    entry_endereco = Text(janela_relatorio, height=1, width=50)
    entry_endereco.pack()

    Label(janela_relatorio, text="Município:").pack(pady=5)
    entry_municipio = Text(janela_relatorio, height=1, width=50)
    entry_municipio.pack()

    Label(janela_relatorio, text="Data de Atualização (dd/mm/aaaa):").pack(pady=5)
    entry_data = Text(janela_relatorio, height=1, width=50)
    entry_data.pack()

    # Campo para detalhar informações
    Label(janela_relatorio, text="Descrição detalhada:").pack(pady=5)
    detalhes_frame = Frame(janela_relatorio)
    detalhes_frame.pack(expand=True, fill=BOTH, pady=10)

    text_detalhes = ScrolledText(detalhes_frame, wrap="word", width=70, height=20)
    text_detalhes.pack(expand=True, fill=BOTH)

    # Botão para gerar o relatório
    def acionar_gerar_relatorio():
        nome_pop = entry_nome_pop.get("1.0", "end").strip()
        endereco = entry_endereco.get("1.0", "end").strip()
        municipio = entry_municipio.get("1.0", "end").strip()
        data_atualizacao = entry_data.get("1.0", "end").strip()
        detalhes = text_detalhes.get("1.0", "end").strip()
        gerar_relatorio(nome_pop, endereco, municipio, data_atualizacao, detalhes)
        registrar_atualizacao(f'Relatório gerado para: {nome_pop}')

    Button(janela_relatorio, text="Gerar Relatório", command=acionar_gerar_relatorio).pack(pady=10)

    

# atualizar as planilhas
def nome_RTP(pop):
    nomeRTP = pop.strip().upper()
    if nomeRTP.startswith('POP-'):
        nomeRTP = nomeRTP.replace('POP-','',1).strip()
        return nomeRTP
def atualizar_dados():
    planilha_forms_caminho = selecionar_arquivo('Selecione a planilha gerada pelo Forms')
    if not planilha_forms_caminho:
        messagebox.showerror("Erro", "Nenhuma planilha do Forms foi selecionada. Encerrando o processo.")
        return

    base_dados_caminho = selecionar_arquivo('Selecione a Base de Dados')
    if not base_dados_caminho:
        messagebox.showerror("Erro", "Nenhuma Base de Dados foi selecionada. Encerrando o processo.")
        return

    planilha_forms = pd.read_excel(planilha_forms_caminho)
    base_dados = pd.read_excel(base_dados_caminho, sheet_name='Panorama POPS RS')

    planilha_forms.columns = planilha_forms.columns.str.strip()
    base_dados.columns = base_dados.columns.str.strip()

    if 'Nome POP' not in planilha_forms.columns or 'Nome POP' not in base_dados.columns:
        messagebox.showerror("Erro", "A coluna 'Nome POP' não foi encontrada em uma das planilhas.")
        return

    planilha_forms['Nome POP'] = planilha_forms['Nome POP'].apply(nome_RTP)
    base_dados['Nome POP'] = base_dados['Nome POP'].apply(nome_RTP)

    for index, row in planilha_forms.iterrows():
        nome_pop = row['Nome POP']
        if nome_pop in base_dados['Nome POP'].values:
            for coluna in planilha_forms.columns:
                if coluna in base_dados.columns:
                    base_dados.loc[base_dados['Nome POP'] == nome_pop, coluna] = row[coluna]
        else:
            base_dados = pd.concat([base_dados, pd.DataFrame([row])], ignore_index=True)

    try:
        workbook = load_workbook(base_dados_caminho)
        sheet = workbook['Panorama POPS RS']
        for i, row in base_dados.iterrows():
            for j, value in enumerate(row):
                column_letter = get_column_letter(j + 1)
                sheet[f"{column_letter}{i + 2}"].value = value

        caminho_salvar = filedialog.asksaveasfilename(defaultextension='.xlsx', filetypes=[('Arquivos Excel', '*.xlsx')])
        if not caminho_salvar:
            messagebox.showerror("Erro", "Nenhum local de salvamento foi selecionado. Encerrando o processo.")
        else:
            workbook.save(caminho_salvar)
            messagebox.showinfo("Sucesso", 'Dados salvos na base DIF/PIR, Obrigado!')
            registrar_atualizacao("Atualização realizada com sucesso")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao salvar a planilha: {e}")

    salvar_onedrive = messagebox.askyesno("Salvar no OneDrive", "Deseja salvar a base de dados atualizada no OneDrive?")
    if salvar_onedrive:
        salvar_no_onedrive()
        acessar_one_drive = messagebox.askyesno('Acesso', 'Deseja acessar o One Drive? (Link Externo)')
        if acessar_one_drive:
            webbrowser.open('https://rsgovbr-my.sharepoint.com/:f:/g/personal/vicenzo-minossi_procergs_rs_gov_br/EkhrwNGgbGVJpK2HN88KFMgB91m2S9OgHyXHcNVvBSa84w?e=8Z4gkM')


def consultar_chatbot(pergunta, contexto):
    try:
        resposta = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": contexto},
                {"role": "user", "content": pergunta}
            ]
        )
        return resposta.choices[0].message['content']
    except Exception as e:
        messagebox.showerror("Erro ao consultar o chatbot", f"{e}")
        return None

#carregar as planilhas 
def carregar_dados_base():
    base_dados_caminho = selecionar_arquivo("Selecione a Base de Dados")
    if not base_dados_caminho:
        messagebox.showerror("Erro", "Nenhuma Base de Dados foi selecionada. Encerrando o processo.")
        return None

    try:
        #aba 'Panorama POPS RS'
        base_dados = pd.read_excel(base_dados_caminho, sheet_name='Panorama POPS RS')
        messagebox.showinfo("Sucesso", "Base de Dados carregada com sucesso!")
        return base_dados
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao carregar a planilha: {e}")
        return None
def atualizar_dados():
    planilha_forms_caminho = selecionar_arquivo('Selecione a planilha gerada pelo Forms')
    if not planilha_forms_caminho:
        messagebox.showerror("Erro", "Nenhuma planilha do Forms foi selecionada. Encerrando o processo.")
        return

    base_dados_caminho = selecionar_arquivo('Selecione a Base de Dados')
    if not base_dados_caminho:
        messagebox.showerror("Erro", "Nenhuma Base de Dados foi selecionada. Encerrando o processo.")
        return

    planilha_forms = pd.read_excel(planilha_forms_caminho)
    base_dados = pd.read_excel(base_dados_caminho, sheet_name='Panorama POPS RS')

    planilha_forms.columns = planilha_forms.columns.str.strip()
    base_dados.columns = base_dados.columns.str.strip()

    if 'Nome POP' not in planilha_forms.columns or 'Nome POP' not in base_dados.columns:
        messagebox.showerror("Erro", "A coluna 'Nome POP' não foi encontrada em uma das planilhas.")
        return

    planilha_forms['Nome POP'] = planilha_forms['Nome POP'].apply(nome_RTP)
    base_dados['Nome POP'] = base_dados['Nome POP'].apply(nome_RTP)

    for index, row in planilha_forms.iterrows():
        nome_pop = row['Nome POP']
        if nome_pop in base_dados['Nome POP'].values:
            for coluna in planilha_forms.columns:
                if coluna in base_dados.columns:
                    base_dados.loc[base_dados['Nome POP'] == nome_pop, coluna] = row[coluna]
        else:
            base_dados = pd.concat([base_dados, pd.DataFrame([row])], ignore_index=True)

    try:
        workbook = load_workbook(base_dados_caminho)
        sheet = workbook['Panorama POPS RS']
        for i, row in base_dados.iterrows():
            for j, value in enumerate(row):
                column_letter = get_column_letter(j + 1)
                sheet[f"{column_letter}{i + 2}"].value = value

        caminho_salvar = filedialog.asksaveasfilename(defaultextension='.xlsx', filetypes=[('Arquivos Excel', '*.xlsx')])
        if not caminho_salvar:
            messagebox.showerror("Erro", "Nenhum local de salvamento foi selecionado. Encerrando o processo.")
        else:
            workbook.save(caminho_salvar)
            messagebox.showinfo("Sucesso", 'Dados salvos na base DIF/PIR, Obrigado!')
            registrar_atualizacao("Atualização realizada com sucesso")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao salvar a planilha: {e}")

    salvar_onedrive = messagebox.askyesno("Salvar no OneDrive", "Deseja salvar a base de dados atualizada no OneDrive?")
    if salvar_onedrive:
        salvar_no_onedrive()
        acessar_one_drive = messagebox.askyesno('Acesso', 'Deseja acessar o One Drive? (Link Externo)')
        if acessar_one_drive:
            webbrowser.open('https://rsgovbr-my.sharepoint.com/:f:/g/personal/vicenzo-minossi_procergs_rs_gov_br/EkhrwNGgbGVJpK2HN88KFMgB91m2S9OgHyXHcNVvBSa84w?e=8Z4gkM')

# Função para abrir a interface do chatbot (em desenvolvimento)
#def abrir_chatbot():
#    # Carrega os dados do Excel
#    base_dados = carregar_dados_base()
    #if base_dados is None:
        #return
    
    # Cria uma janela para o chatbot
    #janela_chatbot = Toplevel(root)
    #janela_chatbot.title("Chatbot - Consulta de POPs")
    #janela_chatbot.geometry("600x500")
    
    # Frame para exibir mensagens
    #frame_chat = Frame(janela_chatbot)
    #frame_chat.pack(expand=True, fill=BOTH)

    #scrollbar = Scrollbar(frame_chat)
    #scrollbar.pack(side=RIGHT, fill=Y)
    
    #chat_log = Text(frame_chat, wrap="word", yscrollcommand=scrollbar.set)
    #chat_log.pack(expand=True, fill=BOTH)
    #scrollbar.config(command=chat_log.yview)

    # Campo de entrada para perguntas
    #Label(janela_chatbot, text="Digite sua pergunta:").pack(pady=5)
    #pergunta_entry = Text(janela_chatbot, height=3, width=70)
    #pergunta_entry.pack()

    # Função para processar a pergunta e exibir a resposta
#    def enviar_pergunta():
#        pergunta = pergunta_entry.get("1.0", "end").strip()
#        if not pergunta:
#            messagebox.showwarning("Aviso", "Digite uma pergunta antes de enviar.")
#           return
#
#        contexto = (
#            "Você é um assistente que responde perguntas sobre pontos de presença (POPs) "
#            "no estado do Rio Grande do Sul. A base de dados contém as seguintes informações: "
#            f"{', '.join(base_dados.columns)}"
#        )
#        resposta = consultar_chatbot(pergunta, contexto)
#        if resposta:
#            chat_log.insert("end", f"Usuário: {pergunta}\n")
#            chat_log.insert("end", f"Chatbot: {resposta}\n\n")
#            chat_log.see("end")
#        pergunta_entry.delete("1.0", "end")
#
 #   # Botão para enviar perguntas
 #   Button(janela_chatbot, text="Enviar", command=enviar_pergunta).pack(pady=10)


def salvar_no_onedrive():
    if user in nomes_autorizados:
        try:
            caminho_base_dados = selecionar_arquivo("Selecione a base de dados para salvar no OneDrive")
            if not caminho_base_dados:
                messagebox.showerror("Erro", "Nenhum arquivo selecionado.")
                return

            base_dados_destino = os.path.join(pasta_onedrive, os.path.basename(caminho_base_dados))
            os.replace(caminho_base_dados, base_dados_destino)

            messagebox.showinfo("Sucesso", "As planilhas foram salvas no OneDrive com sucesso!")
            registrar_atualizacao("Planilhas salvas no OneDrive")
        except Exception as e:
            salvar_manual = messagebox.askyesno("Erro", 'A pasta do One Drive não foi localizada neste dispositivo. Deseja salvar manualmente?')
            if salvar_manual:
                acessar_dados()
                time.sleep(4)
                confirm = messagebox.askyesno('Salvar','As planilhas foram salvas?')
                if confirm:
                    registrar_atualizacao('Planilhas salvas manualmente no One Drive')
    else:
        messagebox.showerror('Erro', 'Sem autorização para acessar o serviço. Consulte o administrador')

def exibir_sobre(): #janela que mostra as info do programa
    try:
        with open(base_sobre, "r") as file:
            conteudo = file.read()

        janela_sobre = Toplevel(root)
        janela_sobre.title("Sobre o programa")
        janela_sobre.geometry("800x500")

        scrollbar = Scrollbar(janela_sobre)
        scrollbar.pack(side="right", fill="y")

        text_widget = Text(janela_sobre, wrap="word", yscrollcommand=scrollbar.set)
        text_widget.pack(expand=True, fill="both")
        scrollbar.config(command=text_widget.yview)

    
        text_widget.insert("1.0", conteudo)

        # leitura
        text_widget.config(state="disabled")
    except FileNotFoundError:
        messagebox.showerror("Erro", "Arquivo de instruções não encontrado.")

def exibir_patch_log():
    try:
        with open(base_teste, "r",encoding='utf-8') as file:
            conteudo = file.read()

        # Criar uma nova janela para exibir o conteúdo
        janela_sobre = Toplevel(root)
        janela_sobre.title("Últimos Patches")
        janela_sobre.geometry("800x500")

        # Adicionar um widget de texto com barra de rolagem para mostrar o conteúdo
        scrollbar = Scrollbar(janela_sobre)
        scrollbar.pack(side="right", fill="y")

        text_widget = Text(janela_sobre, wrap="word", yscrollcommand=scrollbar.set)
        text_widget.pack(expand=True, fill="both")
        scrollbar.config(command=text_widget.yview)

        # Inserir o conteúdo do arquivo no widget de texto
        text_widget.insert("1.0", conteudo)

        # Tornar o widget de texto apenas leitura
        text_widget.config(state="disabled")
    except FileNotFoundError:
        messagebox.showerror("Erro", "Arquivo de instruções não encontrado.")


def consultar_historico_logs():
    try:
        with open(base_logs, "r") as file:
            conteudo = file.readlines()

        # Criar uma nova janela para exibir o histórico
        janela_historico = Toplevel(root)
        janela_historico.title("Histórico de Atualizações")
        janela_historico.geometry("400x300")

        # Adicionar um widget de texto com barra de rolagem para mostrar o histórico
        scrollbar = Scrollbar(janela_historico)
        scrollbar.pack(side="right", fill="y")

        text_widget = Text(janela_historico, wrap="word", yscrollcommand=scrollbar.set)
        text_widget.pack(expand=True, fill="both")
        scrollbar.config(command=text_widget.yview)

        # Inserir o conteúdo do histórico no widget de texto
        if conteudo:
            text_widget.insert("1.0", "".join(conteudo))
        else:
            text_widget.insert("1.0", "Nenhum histórico de atualizações encontrado.")

        # Tornar o widget de texto apenas leitura
        text_widget.config(state="disabled")
    except FileNotFoundError:
        messagebox.showinfo("Histórico de Atualizações", "Nenhum histórico de atualizações encontrado.")

def acessar_dados():
    resposta = messagebox.askyesno("Redirecionamento", "Você será redirecionado para um link externo. Deseja continuar?")
    
    # Verificar a resposta do usuário
    if resposta:
        # Abrir o link no navegador padrão
        webbrowser.open('https://rsgovbr-my.sharepoint.com/:f:/g/personal/vicenzo-minossi_procergs_rs_gov_br/EkhrwNGgbGVJpK2HN88KFMgB91m2S9OgHyXHcNVvBSa84w?e=8Z4gkM')

def selecionar_arquivo(titulo):
    caminho_arquivo = filedialog.askopenfilename(
        title=titulo,
        filetypes=[('Arquivos Excel', '*.xlsx'), ('Todos os arquivos', '*.*')]
    )
    return caminho_arquivo

#carregar dados
def carregar_base_dados():
    base_dados_caminho = selecionar_arquivo('Selecione a Base de Dados')
    if not base_dados_caminho:
        messagebox.showerror("Erro", "Nenhuma Base de Dados foi selecionada. Encerrando o processo.")
        return None, None

    try:
        # Lendo a planilha
        base_dados = pd.read_excel(base_dados_caminho, sheet_name='Panorama POPS RS')
        messagebox.showinfo("Sucesso", f"Base de Dados carregada: {base_dados_caminho}")
        return base_dados, base_dados_caminho
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao carregar a planilha: {e}")
        return None, None


#mapa interativo
def criar_mapa(dados):
    try:
        # check colunas
        colunas_necessarias = ['Latitude', 'Longitude', 'Nome POP', 'Endereço do POP',
         'Condição de equipamentos nos racks','Limpeza','Refrigeração','Rede elétrica']
        for coluna in colunas_necessarias:
            if coluna not in dados.columns:
                raise ValueError(f"A coluna '{coluna}' está ausente na planilha.")

        # Tratar valores ausentes nas colunas de hover
        hover_cols = ['Endereço do POP', 'Condição de equipamentos nos racks']
        for col in hover_cols:
            dados[col] = dados[col].fillna("Sem dados")

        # Converter para números e limpar dados inválidos
        dados['Latitude'] = pd.to_numeric(dados['Latitude'], errors='coerce')
        dados['Longitude'] = pd.to_numeric(dados['Longitude'], errors='coerce')
        dados = dados.dropna(subset=['Latitude', 'Longitude'])

        if dados.empty:
            raise ValueError("Não há dados válidos para plotar no mapa.")

        # Criar mapa com informações detalhadas
        fig = px.scatter_mapbox(
            dados,
            lat='Latitude',
            lon='Longitude',
            text='Nome POP',
            hover_data={
                'Endereço do POP',
                'Condição de equipamentos nos racks',
                'Limpeza',
                'Refrigeração'
                ,'Rede elétrica'
            },
            zoom=5,
            title='Mapa Interativo dos POPs RGS. Dê dois cliques para isolar cada trace e exbir o card de informações'
        )

        # Estilo do mapa
        fig.update_layout(mapbox_style="open-street-map", height=800)

        # Adicionar conexões entre POPs (linhas)
        for i in range(len(dados) - 1):
            fig.add_trace(go.Scattermapbox(
                mode="lines",
                lon=[dados.iloc[i]['Longitude'], dados.iloc[i + 1]['Longitude']],
                lat=[dados.iloc[i]['Latitude'], dados.iloc[i + 1]['Latitude']],
                line=dict(width=2, color='blue'),
                showlegend=True
            ))

        # Exibir o mapa
        fig.show()
    except Exception as e:
        print(e) 

def inicializar_dashboard():
    # Carregar base de dados
    dados, caminho = carregar_base_dados()
    if dados is None:
        return  # Encerrar se não foi possível carregar os dados
    op = messagebox.askokcancel('Redirecionamento','O plot do mapa será feito no navegador. Continuar?')
    if op:
        messagebox.showinfo('Redirecionamento','Mapa sendo gerado, aguarde...')
        criar_mapa(dados)

def easteregg():
    

    pygame.init()

    PRETO = (0, 0, 0)
    BRANCO = (255, 255, 255)
    VERDE = (0, 255, 0)
    VERMELHO = (255, 0 , 0)
    AZUL = (50, 153, 213)

    LARGURA_TELA = 700
    ALTURA_TELA = 700


    TAMANHO_BLoco = 10


    CLOCK = pygame.time.Clock()
    VEL = 15

    tela = pygame.display.set_mode((LARGURA_TELA, ALTURA_TELA))
    pygame.display.set_caption('Conserte os POPS-PAE')


    fundo = pygame.image.load(fundo_game)  # Substitua 'fundo.png' pelo caminho da sua imagem
    fundo = pygame.transform.scale(fundo, (LARGURA_TELA, ALTURA_TELA))  # Ajusta o tamanho da imagem para a tela

    def desenhar_cobrinha(tamanho_bloco, lista_cobrinha):
        for x in lista_cobrinha:
            pygame.draw.rect(tela, VERDE, [x[0], x[1], tamanho_bloco, tamanho_bloco])



    game_over = False
    game_close = False
    

    x1 = LARGURA_TELA // 2
    y1 = ALTURA_TELA // 2
    x1_mudanca = 0
    y1_mudanca = 0

    cobrinha = []
    comprimento_cobrinha = 1

    comida_x = round(random.randrange(0, LARGURA_TELA - TAMANHO_BLoco) / 10.0) * 10.0
    comida_y = round(random.randrange(0, ALTURA_TELA - TAMANHO_BLoco) / 10.0) * 10.0

    while not game_over:
        while game_close:
            tela.fill(AZUL)
            font = pygame.font.SysFont("comicsans", 25)
            mensagem = font.render("Você foi demitido! Pressione C para jogar", True, VERMELHO)
            tela.blit(mensagem, [LARGURA_TELA / 6, ALTURA_TELA / 3])
            pygame.display.update()

            for evento in pygame.event.get():
                if evento.type == pygame.QUIT:
                    game_over = True
                    game_close = False
                if evento.type == pygame.KEYDOWN:
                    if evento.key == pygame.K_q:
                        game_over = True
                        game_close = False
                    if evento.key == pygame.K_c:
                        easteregg()

        for evento in pygame.event.get():
            if evento.type == pygame.QUIT:
                game_over = True
            if evento.type == pygame.KEYDOWN:
                if evento.key == pygame.K_LEFT:
                    x1_mudanca = -TAMANHO_BLoco
                    y1_mudanca = 0
                elif evento.key == pygame.K_RIGHT:
                    x1_mudanca = TAMANHO_BLoco
                    y1_mudanca = 0
                elif evento.key == pygame.K_UP:
                    y1_mudanca = -TAMANHO_BLoco
                    x1_mudanca = 0
                elif evento.key == pygame.K_DOWN:
                    y1_mudanca = TAMANHO_BLoco
                    x1_mudanca = 0

        if x1 >= LARGURA_TELA or x1 < 0 or y1 >= ALTURA_TELA or y1 < 0:
            game_close = True
        x1 += x1_mudanca
        y1 += y1_mudanca

        tela.blit(fundo, (0, 0))

        pygame.draw.rect(tela, VERMELHO, [comida_x, comida_y, TAMANHO_BLoco, TAMANHO_BLoco])

        cabeca_cobrinha = []
        cabeca_cobrinha.append(x1)
        cabeca_cobrinha.append(y1)
        cobrinha.append(cabeca_cobrinha)

        if len(cobrinha) > comprimento_cobrinha:
            del cobrinha[0]

        for x in cobrinha[:-1]:
            if x == cabeca_cobrinha:
                game_close = True

        desenhar_cobrinha(TAMANHO_BLoco, cobrinha)

        pygame.display.update()


        if x1 == comida_x and y1 == comida_y:
            comida_x = round(random.randrange(0, LARGURA_TELA - TAMANHO_BLoco) / 10.0) * 10.0
            comida_y = round(random.randrange(0, ALTURA_TELA - TAMANHO_BLoco) / 10.0) * 10.0
            comprimento_cobrinha += 1

        CLOCK.tick(VEL)

    pygame.quit()
def nao_quero_trabalhar():
    easteregg()

root = Tk()
root.title("Raven V1.1.0")
root.geometry("600x592")
root.config(bg="royalblue") #menu
logo_imagem = Image.open(logo_menu)
logo_imagem = logo_imagem.resize((250, 50))
logo_tk = ImageTk.PhotoImage(logo_imagem)
Label(root, text=(f'{apr}, {user}'), bg="royalblue", fg="white").pack(pady=5)

logo_label = Label(root, image=logo_tk, bg="royalblue")
logo_label.pack(pady=10)
label_rede = Label(root, text=(f' {REDERS_ip}: Conectado a REDERS'), anchor='w', font=('Arial', 10), padx=10)
label_rede.place(relx=0.0, rely=1.0, anchor='sw')  # 'sw' para alinhar no canto inferior esquerdo
#icons dos botões

def criar_botao(root, texto, comando, icone=None, lado='left', padding='15'):
    if icone is not None:
        botao = Button(root, text=texto, command=comando, image=icone, compound=lado)
    else:
        botao = Button(root, text=texto, command=comando)
    botao.pack(pady=padding)
    return botao

def criar_icone(caminho_icone, size=(20, 20)):
    if not os.path.exists(caminho_icone):
        messagebox.showerror("Erro", f"Ícone não encontrado: {caminho_icone}")
        return None
    imagem = Image.open(caminho_icone)
    imagem = imagem.resize(size)
    return ImageTk.PhotoImage(imagem)

icon_atualizar_tk = criar_icone(icon_atualizardados)
icon_acessar_dados_tk = criar_icone(icon_acessardados)
icon_consultar_logs_tk = criar_icone(icon_consultaratt)
icon_onedrive_tk = criar_icone(icon_onedrive)
icon_sobre_tk = criar_icone(icon_sobre)
icon_patchlogs_tk = criar_icone(icon_teste)
icon_gerar_relatorio_tk = criar_icone(icon_relatorio)
icon_raven_tk = criar_icone(icon_raven)
icon_mapa_tk = criar_icone(icon_mapa)



criar_botao(root, ' Atualizar Dados DIF/PIR', atualizar_dados, icon_atualizar_tk)
criar_botao(root, ' Acessar Dados (One Drive)', acessar_dados, icon_acessar_dados_tk)
criar_botao(root, 'Dashboard POPs', inicializar_dashboard,icon_mapa_tk)
criar_botao(root, ' Histórico de atualizações', consultar_historico_logs,icon_consultar_logs_tk)
criar_botao(root, ' Salvar no One Drive', salvar_no_onedrive, icon_onedrive_tk)
criar_botao(root, ' Pactches de melhoria',exibir_patch_log,icon_patchlogs_tk)
criar_botao(root, 'Gerar relatório detalhado (Word)',abrir_janela_relatorio,icon_gerar_relatorio_tk)
criar_botao(root, ' Sobre', exibir_sobre, icon_sobre_tk)
criar_botao(root, 'Não quero trabalhar ',nao_quero_trabalhar)




#problemas de execução do pyinstaller para o arq. executável
#                        DESCOMENTAR PARA EXECUTAR DIAGNÓSTICO 
#pasta_ravenia = encontrar_pasta_ravenia()
#if pasta_ravenia:
   #print(f"Pasta 'RavenIA' encontrada em: {pasta_ravenia}")
#else:
    #print("A pasta 'RavenIA' não foi encontrada.")

root.mainloop()